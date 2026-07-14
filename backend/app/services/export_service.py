"""Export service: render content to pdf/docx/md/txt/json/csv files.

reportlab (PDF) and python-docx (DOCX) are imported lazily. When they are not
installed, the service falls back to writing a plain-text file with the correct
extension so an export always succeeds.
"""

from __future__ import annotations

import csv
import io
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

from app.core.config import settings
from app.core.logging import get_logger
from app.db.models import ExportRecord
from app.repositories.action_repo import ActionRepository
from app.repositories.export_repo import ExportRepository

logger = get_logger(__name__)


def _slug(title: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", (title or "export").strip().lower()).strip("-")
    return slug or "export"


class ExportService:
    def __init__(self, export_repo: ExportRepository, action_repo: ActionRepository) -> None:
        self.export_repo = export_repo
        self.action_repo = action_repo

    def create(self, fmt: str, content: str, title: Optional[str] = None) -> ExportRecord:
        title = title or "Untitled Export"
        base = f"{_slug(title)}-{uuid.uuid4().hex[:8]}.{fmt}"
        dest = Path(settings.exports_dir) / base
        dest.parent.mkdir(parents=True, exist_ok=True)

        if fmt == "pdf":
            self._write_pdf(dest, title, content)
        elif fmt == "docx":
            self._write_docx(dest, title, content)
        elif fmt == "json":
            self._write_json(dest, title, content)
        elif fmt == "csv":
            self._write_csv(dest, content)
        else:  # md, txt and anything else -> plain text
            dest.write_text(content, encoding="utf-8")

        size = dest.stat().st_size if dest.exists() else 0
        record = ExportRecord(
            filename=dest.name,
            format=fmt,
            title=title,
            path=str(dest),
            size=size,
        )
        record = self.export_repo.add(record)
        self.action_repo.log(action=f"export:{fmt}", module="exports", summary=title)
        return record

    # -- writers --------------------------------------------------------------
    def _write_pdf(self, dest: Path, title: str, content: str) -> None:
        try:
            from reportlab.lib.pagesizes import LETTER  # type: ignore
            from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
            from reportlab.lib.units import inch  # type: ignore
            from reportlab.platypus import (  # type: ignore
                Paragraph,
                SimpleDocTemplate,
                Spacer,
            )
        except Exception:
            logger.warning("reportlab not installed; writing PDF as plain text fallback")
            dest.write_text(f"{title}\n\n{content}", encoding="utf-8")
            return

        try:
            doc = SimpleDocTemplate(
                str(dest),
                pagesize=LETTER,
                topMargin=inch,
                bottomMargin=inch,
                leftMargin=inch,
                rightMargin=inch,
            )
            styles = getSampleStyleSheet()
            story = [Paragraph(_escape(title), styles["Title"]), Spacer(1, 12)]
            for para in content.split("\n"):
                if para.strip():
                    story.append(Paragraph(_escape(para), styles["BodyText"]))
                else:
                    story.append(Spacer(1, 8))
            doc.build(story)
        except Exception as exc:  # pragma: no cover
            logger.warning("PDF generation failed (%s); writing text fallback", exc)
            dest.write_text(f"{title}\n\n{content}", encoding="utf-8")

    def _write_docx(self, dest: Path, title: str, content: str) -> None:
        try:
            import docx  # type: ignore
        except Exception:
            logger.warning("python-docx not installed; writing DOCX as plain text fallback")
            dest.write_text(f"{title}\n\n{content}", encoding="utf-8")
            return
        try:
            document = docx.Document()
            document.add_heading(title, level=0)
            for para in content.split("\n"):
                document.add_paragraph(para)
            document.save(str(dest))
        except Exception as exc:  # pragma: no cover
            logger.warning("DOCX generation failed (%s); writing text fallback", exc)
            dest.write_text(f"{title}\n\n{content}", encoding="utf-8")

    def _write_json(self, dest: Path, title: str, content: str) -> None:
        try:
            parsed = json.loads(content)
            payload = parsed
        except (json.JSONDecodeError, ValueError):
            payload = {
                "title": title,
                "content": content,
                "generated_at": datetime.now(timezone.utc).isoformat(),
            }
        dest.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")

    def _write_csv(self, dest: Path, content: str) -> None:
        # If the content already looks like CSV, write it through; otherwise wrap
        # each line as a single-column row.
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        lines = [ln for ln in content.splitlines() if ln.strip() != ""]
        if lines and all("," in ln for ln in lines[:3]):
            for ln in lines:
                writer.writerow(next(csv.reader([ln])))
        else:
            writer.writerow(["content"])
            for ln in lines:
                writer.writerow([ln])
        dest.write_text(buffer.getvalue(), encoding="utf-8")

    # -- retrieval ------------------------------------------------------------
    def list_exports(self) -> List[ExportRecord]:
        return self.export_repo.all_ordered()

    def get(self, export_id: int) -> Optional[ExportRecord]:
        return self.export_repo.get(export_id)


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )
