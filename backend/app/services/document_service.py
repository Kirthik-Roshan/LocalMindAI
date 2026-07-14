"""Document service: ingestion, text extraction, indexing, analysis.

Heavy parsers (pypdf, python-docx) are imported lazily inside methods and
wrapped in try/except so the server boots even when they are not installed.
"""

from __future__ import annotations

import os
import uuid
from pathlib import Path
from typing import List, Optional, Tuple

from app.core.config import settings
from app.core.logging import get_logger
from app.db.models import Document
from app.repositories.action_repo import ActionRepository
from app.repositories.document_repo import DocumentRepository
from app.services.embedding_service import EmbeddingService
from app.services.ollama_service import OllamaService

logger = get_logger(__name__)

SYSTEM_PROMPT = (
    "You are LocalMind, an expert document analyst running fully offline. "
    "Base your answers only on the provided document content. Be accurate, "
    "structured and concise. If the document does not contain the answer, say so."
)

ANALYZE_PROMPTS = {
    "summary": "Write a clear, concise summary of the following document.",
    "key_points": "Extract the most important key points from the document as a Markdown bullet list.",
    "entities": (
        "Identify the key named entities in the document (people, organizations, "
        "locations, dates, amounts). Group them by type in a Markdown list."
    ),
    "timeline": (
        "Extract a chronological timeline of events, dates and milestones from the "
        "document. Format as a Markdown ordered list from earliest to latest."
    ),
    "executive_summary": (
        "Write a polished executive summary of the document suitable for senior "
        "leadership: 1 short paragraph of context, then 3-5 bullet takeaways, then "
        "a recommended next step."
    ),
    "qa": "Answer the user's question using only the document content.",
}


class DocumentService:
    def __init__(
        self,
        doc_repo: DocumentRepository,
        embedding_service: EmbeddingService,
        ollama: OllamaService,
        action_repo: ActionRepository,
    ) -> None:
        self.doc_repo = doc_repo
        self.embedding_service = embedding_service
        self.ollama = ollama
        self.action_repo = action_repo

    # -- ingestion ------------------------------------------------------------
    def upload(self, filename: str, data: bytes) -> Document:
        ext = (Path(filename).suffix.lower().lstrip(".") or "txt")
        safe_name = f"{uuid.uuid4().hex}_{Path(filename).name}"
        dest = Path(settings.uploads_dir) / safe_name
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_bytes(data)

        text, pages = self._extract_text(dest, ext)

        document = Document(
            name=Path(filename).name,
            type=ext,
            size=len(data),
            pages=pages,
            path=str(dest),
            extracted_text=text,
            chunk_count=0,
        )
        document = self.doc_repo.add(document)

        # Auto-index into the vector store.
        try:
            chunks = self.embedding_service.add_document(
                document.id, text, source=document.name
            )
            document.chunk_count = chunks
            self.doc_repo.update(document)
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning("Auto-index failed for document %s: %s", document.id, exc)

        self.action_repo.log(
            action="document_uploaded", module="documents", summary=document.name
        )
        return document

    def _extract_text(self, path: Path, ext: str) -> Tuple[str, Optional[int]]:
        try:
            if ext == "pdf":
                return self._extract_pdf(path)
            if ext in ("docx",):
                return self._extract_docx(path), None
            if ext in ("txt", "md", "markdown", "text", "csv", "json", "log"):
                return path.read_text(encoding="utf-8", errors="ignore"), None
            # Unknown types: best-effort UTF-8 decode.
            return path.read_text(encoding="utf-8", errors="ignore"), None
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning("Text extraction failed for %s: %s", path, exc)
            return "", None

    def _extract_pdf(self, path: Path) -> Tuple[str, Optional[int]]:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            logger.warning("pypdf not installed; cannot extract PDF text")
            return "", None
        try:
            reader = PdfReader(str(path))
            pages = len(reader.pages)
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return "\n\n".join(parts).strip(), pages
        except Exception as exc:  # pragma: no cover
            logger.warning("PDF parse failed: %s", exc)
            return "", None

    def _extract_docx(self, path: Path) -> str:
        try:
            import docx  # type: ignore
        except Exception:
            logger.warning("python-docx not installed; cannot extract DOCX text")
            return ""
        try:
            document = docx.Document(str(path))
            return "\n".join(p.text for p in document.paragraphs).strip()
        except Exception as exc:  # pragma: no cover
            logger.warning("DOCX parse failed: %s", exc)
            return ""

    # -- retrieval ------------------------------------------------------------
    def list_documents(self) -> List[Document]:
        return self.doc_repo.list()

    def get_document(self, document_id: int) -> Optional[Document]:
        return self.doc_repo.get(document_id)

    def delete_document(self, document_id: int) -> bool:
        document = self.doc_repo.get(document_id)
        if document is None:
            return False
        # Remove file + vectors.
        if document.path and os.path.exists(document.path):
            try:
                os.remove(document.path)
            except OSError:
                pass
        self.embedding_service.delete_document(document_id)
        self.doc_repo.delete(document)
        self.action_repo.log(
            action="document_deleted", module="documents", summary=document.name
        )
        return True

    # -- analysis -------------------------------------------------------------
    def _context_for(self, document: Document, limit: int = 12000) -> str:
        text = document.extracted_text or ""
        return text[:limit]

    def analyze(
        self, document_id: int, task: str, question: Optional[str] = None
    ) -> Optional[str]:
        document = self.doc_repo.get(document_id)
        if document is None:
            return None

        instruction = ANALYZE_PROMPTS.get(task, ANALYZE_PROMPTS["summary"])
        context = self._context_for(document)
        if not context.strip():
            return (
                "No readable text could be extracted from this document, so it "
                "cannot be analyzed. The file may be scanned or image-based."
            )

        if task == "qa":
            q = question or "Provide a general overview of this document."
            prompt = (
                f"{instruction}\n\nQuestion: {q}\n\n"
                f"Document '{document.name}':\n---\n{context}\n---"
            )
        else:
            prompt = (
                f"{instruction}\n\nDocument '{document.name}':\n---\n{context}\n---"
            )

        result = self.ollama.generate(prompt, system=SYSTEM_PROMPT)
        self.action_repo.log(
            action=f"analyze:{task}", module="documents", summary=document.name
        )
        return result

    def compare(self, ids: List[int]) -> str:
        documents = [d for d in (self.doc_repo.get(i) for i in ids) if d is not None]
        if len(documents) < 2:
            return "At least two existing documents are required for comparison."

        sections = []
        for doc in documents:
            sections.append(
                f"### Document: {doc.name}\n{self._context_for(doc, limit=5000)}"
            )
        joined = "\n\n".join(sections)
        prompt = (
            "Compare the following documents. Identify their similarities, key "
            "differences, and any contradictions. Present the analysis with "
            "Markdown headings: Overview, Similarities, Differences, Conclusion.\n\n"
            + joined
        )
        result = self.ollama.generate(prompt, system=SYSTEM_PROMPT)
        self.action_repo.log(
            action="document_compare",
            module="documents",
            summary=", ".join(d.name for d in documents),
        )
        return result
